# -*- coding: utf-8 -*-
import re
import os
from six import string_types, text_type, PY2
from docxtpl import DocxTemplate, R, InlineImage, RichText, Listing, Document, Subdoc
from docx.shared import Mm, Inches, Pt
import docx.opc.constants
from docassemble.base.functions import server, this_thread, package_template_filename, get_config
import docassemble.base.filter
from xml.sax.saxutils import escape as html_escape
from docassemble.base.logger import logmessage
from bs4 import BeautifulSoup, NavigableString, Tag
from collections import deque
import PyPDF2
import codecs
from io import open

NoneType = type(None)

DEFAULT_PAGE_WIDTH = '6.5in'

def image_for_docx(fileref, question, tpl, width=None):
    if fileref.__class__.__name__ in ('DAFile', 'DAFileList', 'DAFileCollection', 'DALocalFile'):
        file_info = dict(fullpath=fileref.path())
    else:
        file_info = server.file_finder(fileref, convert={'svg': 'png'}, question=question)
    if 'fullpath' not in file_info:
        return '[FILE NOT FOUND]'
    if width is not None:
        m = re.search(r'^([0-9\.]+) *([A-Za-z]*)', str(width))
        if m:
            amount = float(m.group(1))
            units = m.group(2).lower()
            if units in ['in', 'inches', 'inch']:
                the_width = Inches(amount)
            elif units in ['pt', 'pts', 'point', 'points']:
                the_width = Pt(amount)
            elif units in ['mm', 'millimeter', 'millimeters']:
                the_width = Mm(amount)
            else:
                the_width = Pt(amount)
        else:
            the_width = Inches(2)
    else:
        the_width = Inches(2)
    return InlineImage(tpl, file_info['fullpath'], the_width)

def transform_for_docx(text, question, tpl, width=None):
    if type(text) in (int, float, bool, NoneType):
        return text
    text = text_type(text)
    m = re.search(r'\[FILE ([^,\]]+), *([0-9\.]) *([A-Za-z]+) *\]', text)
    if m:
        amount = m.group(2)
        units = m.group(3).lower()
        if units in ['in', 'inches', 'inch']:
            the_width = Inches(amount)
        elif units in ['pt', 'pts', 'point', 'points']:
            the_width = Pt(amount)
        elif units in ['mm', 'millimeter', 'millimeters']:
            the_width = Mm(amount)
        else:
            the_width = Pt(amount)
        file_info = server.file_finder(m.group(1), convert={'svg': 'png'}, question=question)
        if 'fullpath' not in file_info:
            return '[FILE NOT FOUND]'
        return InlineImage(tpl, file_info['fullpath'], the_width)
    m = re.search(r'\[FILE ([^,\]]+)\]', text)
    if m:
        file_info = server.file_finder(m.group(1), convert={'svg': 'png'}, question=question)
        if 'fullpath' not in file_info:
            return '[FILE NOT FOUND]'
        return InlineImage(tpl, file_info['fullpath'], Inches(2))
    return docassemble.base.filter.docx_template_filter(text, question=question)

def create_hyperlink(url, anchor_text, tpl):
    return InlineHyperlink(tpl, url, anchor_text)

class InlineHyperlink(object):
    def __init__(self, tpl, url, anchor_text):
        self.tpl = tpl
        self.url = url
        self.anchor_text = anchor_text
    def _insert_link(self):
        ref = self.tpl.docx._part.relate_to(self.url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        return '</w:t></w:r><w:hyperlink r:id="%s"><w:r><w:rPr><w:rStyle w:val="InternetLink"/></w:rPr><w:t>%s</w:t></w:r></w:hyperlink><w:r><w:rPr></w:rPr><w:t xml:space="preserve">' % (ref, html_escape(self.anchor_text))
    def __unicode__(self):
        return self._insert_link()
    def __str__(self):
        return self._insert_link()

def include_docx_template(template_file, **kwargs):
    """Include the contents of one docx file inside another docx file."""
    if this_thread.evaluation_context is None:
        return 'ERROR: not in a docx file'
    if template_file.__class__.__name__ in ('DAFile', 'DAFileList', 'DAFileCollection', 'DALocalFile'):
        template_path = template_file.path()
    else:
        template_path = package_template_filename(template_file, package=this_thread.current_package)
    sd = this_thread.misc['docx_template'].new_subdoc()
    sd.subdocx = Document(template_path)
    sd.subdocx._part = sd.docx._part
    first_paragraph = sd.subdocx.paragraphs[0]
    for key, val in kwargs.items():
        if hasattr(val, 'instanceName'):
            the_repr = val.instanceName
        else:
            the_repr = '_codecs.decode(_array.array("b", "' + re.sub(r'\n', '', codecs.encode(bytearray(val, encoding='utf-8'), 'base64').decode()) + '".encode()), "base64").decode()'
        first_paragraph.insert_paragraph_before(str("{%%p set %s = %s %%}" % (key, the_repr)))
    if 'docx_include_count' not in this_thread.misc:
        this_thread.misc['docx_include_count'] = 0
    this_thread.misc['docx_include_count'] += 1
    return sd

def add_to_rt(tpl, rt, parsed):
    list_tab = False
    block_tab = False
    ordered_list = 0
    while (len(list(parsed)) > 0):
        html_names =    {
            'em': False,
            'code': False,
            'strong': False,
            'h1': False,
            'h2': False,
            'h3': False,
            'h4': False,
            'u': False,
            'a': False,
            'strike': False,
            'ol': False,
            'ul': False,
            'li': False,
            'blockquote': False
        }
        href = ''
        html_out = parsed.popleft()
        for parent in html_out.parents:
            for html_key, html_value in html_names.items():
                if (parent.name ==  html_key):
                    html_names[html_key] = True
                    if (html_key == 'a'):
                        href = parent.get('href')
        rtf_pretext = ''
        if (html_names['code']):
            html_names['em'] = True
        if (html_names['ol'] or html_names['ul']):
            if (html_out == '\n'):
                list_tab = True
            elif (list_tab == True):
                if (html_names['ol']):
                    ordered_list += 1
                    rt.add('\t' + str(ordered_list) + '. ')
                else:
                    rt.add('\t- ')
                list_tab = False
        else:
            list_tab = False
        if (html_names['blockquote']):
            if (html_out == '\n'):
                block_tab = True
            elif (block_tab == True):
                rt.add('\t')
                block_tab = False
        else:
            block_tab = False
        if (html_names['a']):
            rt.add(rtf_pretext + html_out, italic=html_names['em'],
                bold=html_names['strong'], underline=True, strike=html_names['strike'],
                url_id=tpl.build_url_id(href))
        elif (html_names['h1']):
            if (html_names['a']):
                rt.add(rtf_pretext + html_out, italic=html_names['em'],
                    bold=True, underline=True, strike=html_names['strike'],
                    url_id=tpl.build_url_id(href), size=60)
            else:
                rt.add(rtf_pretext + html_out, italic=html_names['em'],
                    bold=True, underline=html_names['u'], strike=html_names['strike'], size=60)
        elif (html_names['h2']):
            if (html_names['a']):
                rt.add(rtf_pretext + html_out, italic=html_names['em'],
                    bold=True, underline=True, strike=html_names['strike'],
                    url_id=tpl.build_url_id(href), size=40)
            else:
                rt.add(rtf_pretext + html_out, italic=html_names['em'],
                    bold=True, underline=html_names['u'], strike=html_names['strike'], size=40)
        elif (html_names['h3']):
            if (html_names['a']):
                rt.add(rtf_pretext + html_out, italic=html_names['em'],
                    bold=True, underline=True, strike=html_names['strike'],
                    url_id=tpl.build_url_id(href), size=30)
            else:
                rt.add(rtf_pretext + html_out, italic=html_names['em'],
                    bold=True, underline=html_names['u'], strike=html_names['strike'], size=30)
        elif (html_names['h4']):
            if (html_names['a']):
                rt.add(rtf_pretext + html_out, italic=html_names['em'],
                    bold=True, underline=True, strike=html_names['strike'],
                    url_id=tpl.build_url_id(href), size=20)
            else:
                rt.add(rtf_pretext + html_out, italic=html_names['em'],
                    bold=True, underline=html_names['u'], strike=html_names['strike'], size=20)
        else:
            rt.add(rtf_pretext + html_out, italic=html_names['em'],
                bold=html_names['strong'], underline=html_names['u'], strike=html_names['strike'])
    return rt

def get_children(descendants, parsed):
    subelement = False
    descendants_buff = deque()
    if descendants is None:
        return descendants_buff
    if (isinstance(descendants, NavigableString)):
        parsed.append(descendants)
    else:
        for child in descendants.children:
            if (child.name == None):
                if (subelement == False):
                    parsed.append(child)
                else:
                    descendants_buff.append(child)
            else:
                if (subelement == False):
                    subelement = True
                    descendants_buff.append(child)
                else:
                    descendants_buff.append(child)
    descendants_buff.reverse()
    return descendants_buff

def html_linear_parse(soup):
    html_tag = soup.html
    descendants = deque()
    descendants.appendleft(html_tag)
    parsed = deque()
    while (len(list(descendants)) > 0):
        child = descendants.popleft()
        from_children = get_children(child, parsed)
        descendants.extendleft(from_children)
    return parsed

class SoupParser(object):
    def __init__(self, tpl):
        self.paragraphs = [dict(params=dict(style='p', indentation=0), runs=[RichText('')])]
        self.current_paragraph = self.paragraphs[-1]
        self.run = self.current_paragraph['runs'][-1]
        self.bold = False
        self.italic = False
        self.underline = False
        self.strike = False
        self.indentation = 0
        self.style = 'p'
        self.still_new = True
        self.size = None
        self.tpl = tpl
    def new_paragraph(self):
        if self.still_new:
            logmessage("new_paragraph is still new and style is " + self.style + " and indentation is " + text_type(self.indentation))
            self.current_paragraph['params']['style'] = self.style
            self.current_paragraph['params']['indentation'] = self.indentation
            return
        logmessage("new_paragraph where style is " + self.style + " and indentation is " + text_type(self.indentation))
        self.current_paragraph = dict(params=dict(style=self.style, indentation=self.indentation), runs=[RichText('')])
        self.paragraphs.append(self.current_paragraph)
        self.run = self.current_paragraph['runs'][-1]
        self.still_new = True
    def __str__(self):
        return self.__unicode__().encode('utf-8') if PY2 else self.__unicode__()
    def __unicode__(self):
        output = ''
        list_number = 1
        for para in self.paragraphs:
            logmessage("Got a paragraph where style is " + para['params']['style'] + " and indentation is " + text_type(para['params']['indentation']))
            output += '<w:p><w:pPr><w:pStyle w:val="Normal"/>'
            if para['params']['style'] in ('ul', 'ol', 'blockquote'):
                output += '<w:ind w:left="' + text_type(36*para['params']['indentation']) + '" w:right="0" w:hanging="0"/>'
            output += '<w:rPr></w:rPr></w:pPr>'
            if para['params']['style'] == 'ul':
                output += text_type(RichText("•\t"))
            if para['params']['style'] == 'ol':
                output += text_type(RichText(text_type(list_number) + ".\t"))
                list_number += 1
            else:
                list_number = 1
            for run in para['runs']:
                output += text_type(run)
            output += '</w:p>'
        return output
    def start_link(self, url):
        ref = self.tpl.docx._part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        self.current_paragraph['runs'].append('<w:hyperlink r:id="%s">' % (ref, ))
        self.new_run()
        self.still_new = False
    def end_link(self):
        self.current_paragraph['runs'].append('</w:hyperlink>')
        self.new_run()
        self.still_new = False
    def new_run(self):
        self.current_paragraph['runs'].append(RichText(''))
        self.run = self.current_paragraph['runs'][-1]
    def traverse(self, elem):
        for part in elem.contents:
            if isinstance(part, NavigableString):
                self.run.add(text_type(part), italic=self.italic, bold=self.bold, underline=self.underline, strike=self.strike, size=self.size)
                self.still_new = False
            elif isinstance(part, Tag):
                logmessage("Part name is " + text_type(part.name))
                if part.name == 'p':
                    self.new_paragraph()
                    self.traverse(part)
                elif part.name == 'li':
                    self.new_paragraph()
                    self.traverse(part)
                elif part.name == 'ul':
                    logmessage("Entering a UL")
                    oldstyle = self.style
                    self.style = 'ul'
                    self.indentation += 10
                    self.traverse(part)
                    self.indentation -= 10
                    self.style = oldstyle
                    logmessage("Leaving a UL")
                elif part.name == 'ol':
                    logmessage("Entering a OL")
                    oldstyle = self.style
                    self.style = 'ol'
                    self.indentation += 10
                    self.traverse(part)
                    self.indentation -= 10
                    self.style = oldstyle
                    logmessage("Leaving a OL")
                elif part.name == 'strong':
                    self.bold = True
                    self.traverse(part)
                    self.bold = False
                elif part.name == 'em':
                    self.italic = True
                    self.traverse(part)
                    self.italic = False
                elif part.name == 'strike':
                    self.strike = True
                    self.traverse(part)
                    self.strike = False
                elif part.name == 'u':
                    self.underline = True
                    self.traverse(part)
                    self.underline = False
                elif part.name == 'blockquote':
                    oldstyle = self.style
                    self.style = 'blockquote'
                    self.indentation += 20
                    self.traverse(part)
                    self.indentation -= 20
                    self.style = oldstyle
                elif re.match(r'h[1-6]', part.name):
                    oldsize = self.size
                    self.size = 60 - ((int(part.name[1]) - 1) * 10)
                    self.new_paragraph()
                    self.bold = True
                    self.traverse(part)
                    self.bold = False
                    self.size = oldsize
                elif part.name == 'a':
                    self.start_link(part['href'])
                    self.underline = True
                    self.traverse(part)
                    self.underline = False
                    self.end_link()
            else:
                logmessage("Encountered a " + part.__class__.__name__)

def markdown_to_docx(text, tpl):
    if get_config('new markdown to docx', False):
        source_code = docassemble.base.filter.markdown_to_html(text, do_terms=False)
        source_code = re.sub("\n", ' ', source_code)
        source_code = re.sub(">\s+<", '><', source_code)
        soup = BeautifulSoup('<html>' + source_code + '</html>', 'html.parser')
        parser = SoupParser(tpl)
        for elem in soup.find_all(recursive=False):
            parser.traverse(elem)
        output = text_type(parser)
        logmessage(output)
        return output
    else:
        source_code = docassemble.base.filter.markdown_to_html(text, do_terms=False)
        source_code = re.sub(r'(?<!\>)\n', ' ', source_code)
        #source_code = re.sub("\n", ' ', source_code)
        #source_code = re.sub(">\s+<", '><', source_code)
        rt = RichText('')
        soup = BeautifulSoup(source_code, 'lxml')
        html_parsed = deque()
        html_parsed = html_linear_parse(soup)
        rt = add_to_rt(tpl, rt, html_parsed)
        return rt

def pdf_pages(file_info, width):
    output = ''
    if width is None:
        width = DEFAULT_PAGE_WIDTH
    if not os.path.isfile(file_info['path'] + '.pdf'):
        if file_info['extension'] in ('rtf', 'doc', 'odt') and not os.path.isfile(file_info['path'] + '.pdf'):
            server.fg_make_pdf_for_word_path(file_info['path'], file_info['extension'])
    if 'pages' not in file_info:
        try:
            reader = PyPDF2.PdfFileReader(open(file_info['path'] + '.pdf', 'rb'))
            file_info['pages'] = reader.getNumPages()
        except:
            file_info['pages'] = 1
    max_pages = 1 + int(file_info['pages'])
    formatter = '%0' + text_type(len(text_type(max_pages))) + 'd'
    for page in range(1, max_pages):
        page_file = dict()
        test_path = file_info['path'] + 'page-in-progress'
        if os.path.isfile(test_path):
            while (os.path.isfile(test_path) and time.time() - os.stat(test_path)[stat.ST_MTIME]) < 30:
                if not os.path.isfile(test_path):
                    break
                time.sleep(1)
        page_file['extension'] = 'png'
        page_file['path'] = file_info['path'] + 'page-' + formatter % page
        page_file['fullpath'] = page_file['path'] + '.png'
        if not os.path.isfile(page_file['fullpath']):
            server.fg_make_png_for_pdf_path(file_info['path'] + '.pdf', 'page')
        if os.path.isfile(page_file['fullpath']):
            output += text_type(image_for_docx(docassemble.base.functions.DALocalFile(page_file['fullpath']), docassemble.base.functions.this_thread.current_question, docassemble.base.functions.this_thread.misc.get('docx_template', None), width=width))
        else:
            output += "[Error including page image]"
        output += ' '
    return(output)
